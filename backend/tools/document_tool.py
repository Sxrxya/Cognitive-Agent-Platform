"""
Cognitive Agent Platform — Document Tool
Document ingestion, parsing, and summarization.
"""

import structlog
from services.llm_service import get_llm_service

logger = structlog.get_logger(__name__)


def summarize_text(text: str, max_length: int = 500) -> str:
    """Summarize a text document using LLM."""
    llm = get_llm_service()

    # Truncate input if too long for a single call
    if len(text) > 15000:
        text = text[:15000] + "\n\n[... content truncated for summarization]"

    result = llm.generate(
        f"Summarize the following text in {max_length} words or fewer:\n\n{text}",
        system_instruction=(
            "You are a precise summarizer. Create a clear, structured summary "
            "capturing the key points, conclusions, and important details."
        ),
    )
    logger.info("Text summarized", input_len=len(text), output_len=len(result))
    return result


def extract_key_topics(text: str) -> list[str]:
    """Extract key topics from text."""
    llm = get_llm_service()

    if len(text) > 10000:
        text = text[:10000]

    result = llm.generate(
        f"Extract the top 5-10 key topics from the following text. Return as a JSON array of strings only:\n\n{text}",
        system_instruction="You are a topic extraction engine. Return only a JSON array of topic strings, nothing else.",
    )

    # Try to parse as JSON
    import json
    try:
        start = result.find("[")
        end = result.rfind("]") + 1
        if start >= 0 and end > start:
            return json.loads(result[start:end])
    except json.JSONDecodeError:
        pass

    # Fallback: split by newlines
    return [line.strip("- ").strip() for line in result.strip().splitlines() if line.strip()]


def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        logger.info("PDF parsed", path=file_path, pages=len(reader.pages), chars=len(text))
        return text
    except Exception as e:
        logger.error("PDF parsing failed", path=file_path, error=str(e))
        return f"Failed to parse PDF: {str(e)}"


def parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(file_path)
        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        logger.info("DOCX parsed", path=file_path, paragraphs=len(doc.paragraphs), chars=len(text))
        return text
    except Exception as e:
        logger.error("DOCX parsing failed", path=file_path, error=str(e))
        return f"Failed to parse DOCX: {str(e)}"
